from io import BytesIO

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


def build_generated_document_docx(
    *,
    title: str,
    content: str,
    language: str,
    disclaimer: str | None = None,
) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_paragraph.add_run(
        "NorthBridgeAI"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)

    doc.add_paragraph("")

    for block in (content or "").split("\n\n"):
        cleaned = block.strip()
        if not cleaned:
            continue

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for line_index, line in enumerate(cleaned.splitlines()):
            run = p.add_run(line.strip())
            run.font.size = Pt(11)
            if line_index < len(cleaned.splitlines()) - 1:
                p.add_run("\n")

    if disclaimer:
        doc.add_paragraph("")
        note = doc.add_paragraph()
        note_run = note.add_run(disclaimer)
        note_run.italic = True
        note_run.font.size = Pt(9)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()